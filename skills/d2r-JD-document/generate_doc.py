#!/usr/bin/env python3
"""Generate a .docx conforming to the corporate template (font, structure, footer).

Usage:
    python generate_doc.py --data data.json [--config config.json] --output out.docx

data.json shape:
{
  "title": "Q3 Regional Sales Review",
  "subtitle": "Prepared for the Leadership Team",   # optional
  "author": "Jane Doe",                             # optional
  "date": "August 18, 2026",                         # optional
  "sections": [
    {"heading": "Executive Summary", "level": 1, "body": ["Paragraph one.", "Paragraph two."]},
    {"heading": "Background",        "level": 1, "body": ["..."]},
    {"heading": "Details",           "level": 2, "body": ["..."]}
  ]
}

config.json (company-specific, all fields optional, shown values are defaults):
{
  "company_name": "Your Company Name",
  "font_name": "Calibri",
  "body_size_pt": 11,
  "heading_sizes_pt": {"1": 16, "2": 13, "3": 12},
  "heading_color_hex": "1F4E79",
  "footer_text": "{company_name}",
  "page_size": "Letter",
  "logo_path": null
}
"""
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULT_CONFIG = {
    "company_name": "Your Company Name",
    "font_name": "Calibri",
    "body_size_pt": 11,
    "heading_sizes_pt": {"1": 16, "2": 13, "3": 12},
    "heading_color_hex": "1F4E79",
    "footer_text": "{company_name}",
    "page_size": "Letter",
    "logo_path": None,
}


def load_config(path):
    cfg = dict(DEFAULT_CONFIG)
    if path:
        with open(path, "r", encoding="utf-8") as f:
            user_cfg = json.load(f)
        cfg.update(user_cfg)
        if "heading_sizes_pt" in user_cfg:
            merged = dict(DEFAULT_CONFIG["heading_sizes_pt"])
            merged.update(user_cfg["heading_sizes_pt"])
            cfg["heading_sizes_pt"] = merged
        if cfg.get("logo_path") and not Path(cfg["logo_path"]).is_absolute():
            cfg["logo_path"] = str((Path(path).parent / cfg["logo_path"]).resolve())
    return cfg


def set_base_fonts(doc, cfg):
    """Force Calibri (or configured font) across Normal, Title, and Heading styles."""
    font_name = cfg["font_name"]
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:cs"), font_name)

    doc.styles["Normal"].font.size = Pt(cfg["body_size_pt"])

    color = RGBColor.from_string(cfg["heading_color_hex"])
    for level_str, size in cfg["heading_sizes_pt"].items():
        style = doc.styles[f"Heading {level_str}"]
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def set_page_size_and_margins(section, cfg):
    if cfg["page_size"].lower() == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_footer(section, cfg):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = cfg["footer_text"].format(company_name=cfg["company_name"])
    run = p.add_run(f"{text}  |  Page ")
    run.font.name = cfg["font_name"]
    run.font.size = Pt(9)
    add_page_number_field(p)


def add_title_page(doc, cfg, data):
    if cfg.get("logo_path"):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(cfg["logo_path"], width=Inches(1.5))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(data["title"])

    if data.get("subtitle"):
        sub_p = doc.add_paragraph(style="Subtitle")
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.add_run(data["subtitle"])

    doc.add_paragraph()
    doc.add_paragraph()

    if data.get("author"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(data["author"])

    if data.get("date"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(data["date"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(cfg["company_name"]).font.size = Pt(10)

    doc.add_page_break()


def add_body(doc, data, cfg):
    for section_data in data.get("sections", []):
        level = section_data.get("level", 1)
        doc.add_heading(section_data["heading"], level=level)
        for para_text in section_data.get("body", []):
            doc.add_paragraph(para_text)


def build_document(data, cfg):
    doc = Document()
    set_base_fonts(doc, cfg)
    set_page_size_and_margins(doc.sections[0], cfg)
    add_footer(doc.sections[0], cfg)
    add_title_page(doc, cfg, data)
    add_body(doc, data, cfg)
    return doc


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data", required=True, help="Path to data.json describing the document content")
    parser.add_argument("--config", default=None, help="Path to config.json with company font/branding overrides")
    parser.add_argument("--output", required=True, help="Path to write the generated .docx")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)
    cfg = load_config(args.config)

    doc = build_document(data, cfg)
    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    sys.exit(main())
